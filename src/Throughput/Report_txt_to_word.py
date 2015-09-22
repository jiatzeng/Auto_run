# -*- coding: utf8 -*-

'''
Created on 2010/1/11

@author: Jocabion
'''


import time
import re
import subprocess
import getpass
import sys
import telnetlib
import SendKeys
import socket
import random
import string   
import logging
import os,sys
import xlrd 
import fileinput
import datetime
import shutil, errno
import subprocess
import distutils.dir_util
import wmi 
import math
import _winreg
import ctypes



from threading import Thread


from datetime import datetime
from ConfigParser import SafeConfigParser
from ConfigParser import RawConfigParser
from docx import Document
from docx.shared import Inches


from subprocess import Popen,PIPE
from decimal import Decimal


if __name__ == '__main__':
    
    
    save_profile = os.getcwd() + '//Profile'
    
    save_result = os.getcwd() +'//Result//Throughput2g.txt'
    
    chrome_path = os.getcwd() + '//chromedriver_win32//chromedriver.exe'
    
    chariot_path = os.getcwd() + '/Chariot_tst'       
              
    document = Document(os.getcwd() +'\jia.docx')
    
    
 
    file = open(save_result, 'r')
    
    
    lines = file.readlines()
    print lines
    
    list_open=[]
    
    list_wep64 = []
    
    list_wep128 = []
    
    list_tkip = []
    
    list_aes = []
    
    list_channel = []
    
    list_initial = range(21)
    
    
    for item in lines :
        if  item.find('TX')!= -1 or item.find('RX')!= -1 :
#           print item
          if 'Open' in item :
              item = re.findall('[0-9]{1,3}',str(item))
              print item[0]
              list_open.append(item[0])
          elif 'WEP_64' in item :    
              item = re.findall('[0-9]{1,3}',str(item))
              print item[0]
              list_wep64.append(item[0])
          elif 'WEP_128' in item :    
              item = re.findall('[0-9]{1,3}',str(item))
              print item[0]    
              list_wep128.append(item[0])
          elif 'TKIP' in item :    
              item = re.findall('[0-9]{1,3}',str(item))
              print item[0]  
              list_tkip.append(item[0])   
          elif 'AES' in item :    
              item = re.findall('[0-9]{1,3}',str(item))
              print item[0]    
              list_aes.append(item[0])   
          elif 'Channel' in item :    
              item = re.findall('[0-9]{1,3}',str(item))
              print item[1]  
              list_channel.append(item[1])   
              
        else :
           pass 

    

    file.close()
        

    print  list_open
    
    print list_wep64
    
    print list_wep128
    
    print list_tkip
    
    print list_aes
    
    print list_channel
    
    
    
    list_all =  list_initial + list_open + ['WEP-64']+list_wep64 + ['WEP-128']+ list_wep128 +['WPA-TKIP']+list_tkip +['WPA2-AES']+ list_aes      
    
    print list_all
    
    
    
    
    
    list =[]
    
    
    for table in document.tables:
     for row in table.rows:
        for cell in row.cells:
#             match_text = re.split('\n',cell.text)
            
#             print cell.text
            
            list.append(cell.text)
            
#     print list
#     sum = 0
#     for item in list :
#          
#         print sum
#         print item
#         sum = sum + 1
    
    for i in range(0,len(list),1):
      if i >20 and i < 70 :
        
        print list_all[i]
        list[i] = list_all[i]
        
#       elif i > 30 and i < 40 :
#           
#         print list[i]   
#       
#       elif i > 40 and i < 50 :
#           
#         print list[i]     
#         
#       elif i > 50 and i < 60 :
#           
#         print list[i]     
#         
#       elif i > 60 and i < 70 :
#           
#         print list[i]    
#         

#         
#       elif i > 100 and i < 110 :
#           
#         print list[i]     
#         
#       elif i > 110 and i < 120 :
#           
#         print list[i]   
#         
#       elif i > 120 and i < 130 :
#           
#         print list[i]   
#       
#       elif i > 130 and i < 140 :
#           
#         print list[i]   
#         
#       elif i > 176 and i < 180 :
#           
#         print list[i]    
#       
#       elif i > 180 and i < 184 :
#           
#         print list[i]  
#       
#       elif i > 184 and i < 188 :
#           
#         print list[i]    
#             
#       elif i > 188 and i < 192 :
#           
#         print list[i]    
#       
#       elif i > 192 and i < 196 :
#           
#         print list[i]  
#       
    
        
#     print list
    
    
    
#     for table in document.tables:
#      for row in table.rows:
#        for cell in row.cells:   
    print len(document.tables)
    print len(document.tables[0].rows)
    print len(document.tables[0].rows[0].cells)
    
    sum = 0
    l = 0

    for k in range(0,len(document.tables),1):
     for i in range(0,len(document.tables[l].rows),1):    
      for j in range(0,len(document.tables[l].rows[l].cells),1): 
#            for i in range(0,len(list),1): 
               
               print document.tables[k].rows[i].cells[j].text 
               document.tables[k].rows[i].cells[j].text = list[sum]
               sum = sum + 1
               print sum
  
     l = l+1
  
  
  
  
                 
#     for j in range(0,10,1): 
# 
#                print document.tables[0].rows[1].cells[j].text 
               
               
               
               

#     print row.cells[j].text
    
    
    
    
    
    
    document.save('jia.docx')
        
#         print list.index(str(item))
#         elif item == 'WEP-64' :  
#             print list.index(str(item))  
#         elif item == 'WEP-128' :
#             print list.index(str(item)) 
#         elif item == 'WPA-TKIP' :    
#             print list.index(str(item)) 
#         elif item == 'WPA2-AES' :   
#             print list.index(str(item)) 
 
                
            
            
        
    
            
#             match_text = re.findall('[^a-zA-z]',cell.text)
#             sum =''
#             for i in range(0,len(match_text),1) :
#  
#                sum = sum + str(match_text[i])
#                print sum
    


